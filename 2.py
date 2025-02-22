import docx
import random
import os
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog
import re  # Import the regular expression module

def extract_questions(docx_path):
    try:
        doc = docx.Document(docx_path)
    except Exception as e:
        messagebox.showerror("错误", f"无法打开文档: {e}\n请检查文件路径和格式是否正确。")
        return []

    questions = []
    current_question = {}
    current_type = None
    in_question = False  # are we inside a question

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        # print(f"DEBUG: Processing paragraph: '{text}'")

        # --- Detect Question Type ---
        if text.startswith("一、单项选择题"):
            current_type = "单选题"
            in_question = False
            try:
                num_questions_type = int(text.split("（")[1].split("道")[0])
            except (IndexError, ValueError):
                num_questions_type = 0
            current_question = {"题型": current_type, "数量": num_questions_type}
            continue

        elif text.startswith("二、多项选择题"):
            current_type = "多选题"
            in_question = False
            try:
                num_questions_type = int(text.split("（")[1].split("道")[0])
            except (IndexError, ValueError):
                num_questions_type = 0
            current_question = {"题型": current_type, "数量": num_questions_type}
            continue

        elif text.startswith("三、判断题"):
            current_type = "判断题"
            in_question = False
            try:
                num_questions_type = int(text.split("（")[1].split("道")[0])
            except (IndexError, ValueError):
                num_questions_type = 0
            current_question = {"题型": current_type, "数量": num_questions_type}
            continue

        elif text.startswith("四、简答题"):
            current_type = "简答题"
            in_question = False
            try:
                num_questions_type = int(text.split("（")[1].split("道")[0])
            except (IndexError, ValueError):
                num_questions_type = 0
            current_question = {"题型": current_type, "数量": num_questions_type}
            continue

        if current_type is None:
            continue

        # --- Question Start Detection ---
        if current_type in ("单选题", "多选题", "判断题", "简答题") and text and text[0].isdigit():
            if in_question and current_question and current_question.get("题型") is not None:
                questions.append(current_question)

            in_question = True
            try:
                parts = re.split(r'(\d+)(?:\.\s*|\s+)', text, 1)  # More flexible regex
                if len(parts) > 2:
                    question_number = int(parts[1])
                    question_text = parts[2].strip()

                    # --- Answer Extraction ---
                    answer = ""
                    if current_type in ("单选题", "多选题"):
                        match = re.search(r'[（(](.+?)[)）]', question_text)
                        if match:
                            answer = match.group(1).strip().upper()
                            question_text = re.sub(r'[（(].*?[)）]', '', question_text).strip()
                    elif current_type == "判断题":
                        # --- Robust True/False Answer Extraction ---
                        answer = ""  # Default to empty
                        question_text_lower = question_text.lower()  # For case-insensitive matching

                        # Check for answers in parentheses
                        match = re.search(r'[（(]\s*([√×对错TtFf])\s*[)）]', question_text_lower)
                        if match:
                            answer_char = match.group(1).strip()
                            if answer_char in ('√', '对', 't'):
                                answer = "√"  # Standardize to checkmark
                            elif answer_char in ('×', '错', 'f'):
                                answer = "×"  # Standardize to x
                            question_text = re.sub(r'[（(].*?[)）]', '', question_text).strip()  # Remove (answer)

                        # Check for answers NOT in parentheses (e.g., "正确", "错误")
                        elif "正确" in question_text_lower:
                            answer = "√"
                            question_text = question_text.replace("正确", "").replace("错误", "").strip() #Remove the keyword
                        elif "错误" in question_text_lower:
                            answer = "×"
                            question_text = question_text.replace("正确", "").replace("错误", "").strip()

                    # ---  Short Answer Extraction Here ---
                    elif current_type == "简答题":
                        match = re.search(r'答[:：]\s*(.*)', question_text)  # Handles variations
                        if match:
                            answer = match.group(1).strip()
                            question_text = re.sub(r'答[:：].*', '', question_text).strip()  # Remove answer

                    current_question = {
                        "题型": current_type,
                        "题目": question_text,
                        "答案": answer,
                        "题号": question_number  # Keep original question number
                    }
                else:
                    in_question = False
                    current_question = {}
            except (IndexError, ValueError) as e:
                print(f"    DEBUG: Error processing question start: {e}")
                in_question = False
                current_question = {}
            continue

        # --- Option and Continued Question Text Handling ---
        if in_question:
            # --- *FIRST* Append Text (if any) ---
            if text:
                if "题目" in current_question:
                    current_question["题目"] += " " + text

            # --- *THEN* Check for Options ---
            if current_type in ("单选题", "多选题") and text.startswith(("A.", "B.", "C.", "D.", "E.")):
                try:
                    parts = re.split(r'([A-E]\.)\s*', text, 1)
                    if len(parts) > 2:
                        option_label = parts[1].strip()
                        option_text = parts[2].strip()
                        if option_label not in current_question:
                            current_question[option_label] = option_text
                        else:
                            print(f"    DEBUG: Duplicate option '{option_label}' found.  Keeping the first occurrence.")
                except IndexError as e:
                    print(f"    DEBUG: IndexError during option extraction: {e}")

            # --- Moved Short Answer answer logic to question start
            # elif current_type == "简答题" and text.startswith("答："):
            #     current_question["答案"] = text[2:].strip()


    if in_question and current_question and current_question.get("题型") is not None:
        questions.append(current_question)

    return questions




def create_exam_paper(questions, num_single, num_multi, num_judge, num_short, output_path):
    doc = docx.Document()
    doc.add_heading('考试试卷', level=1)

    # --- Combine all questions into a single list ---
    all_questions = [q for q in questions if q.get("题型") in ("单选题", "多选题", "判断题", "简答题")]

    # --- Calculate total number of questions to select ---
    total_questions_to_select = num_single + num_multi + num_judge + num_short

    # --- Create a list of indices and shuffle it ---
    if len(all_questions) <= total_questions_to_select:
        # If we're selecting ALL questions (or more), just shuffle the all_questions directly.
        selected_questions_all = all_questions
        random.shuffle(selected_questions_all)
    else:
        #  If there are more questions, it is necessary to sample them without replacement.
        indices = list(range(len(all_questions)))  # Create a list of indices [0, 1, 2, ...]
        random.shuffle(indices)  # Shuffle the indices randomly
        selected_indices = indices[:total_questions_to_select]  # Take the first N indices
        selected_questions_all = [all_questions[i] for i in selected_indices]  # Get questions by the indices

    # --- Filter questions by type (from the selected questions!)---
    selected_single = [q for q in selected_questions_all if q["题型"] == "单选题"]
    selected_multi = [q for q in selected_questions_all if q["题型"] == "多选题"]
    selected_judge = [q for q in selected_questions_all if q["题型"] == "判断题"]
    selected_short = [q for q in selected_questions_all if q["题型"] == "简答题"]



     # --- Add single-choice questions ---
    if selected_single:
        doc.add_heading('一、单项选择题', level=2)
        for i, q in enumerate(selected_single):
            # --- Check if '题目' key exists and is not None ---
            if '题目' in q and q['题目'] is not None:
                doc.add_paragraph(f"{q['题号']}. {q['题目']}")  # Use stored question number
            else:
                print(f"DEBUG: Skipping question {i+1} due to missing or empty '题目' key.")  # Debug print
                continue # Skip to the next question

            for option in ["A", "B", "C", "D"]:
                if option in q and q[option] is not None:
                    doc.add_paragraph(f"   {option}. {q[option]}")
            doc.add_paragraph("")


    # --- Add multiple-choice questions ---
    if selected_multi:
        doc.add_heading('二、多项选择题', level=2)
        for i, q in enumerate(selected_multi):
            # ---  Check if '题目' key exists and is not None ---
            if '题目' in q and q['题目'] is not None:
                doc.add_paragraph(f"{q['题号']}. {q['题目']}")  # Use stored question number
            else:
                print(f"DEBUG: Skipping question {i+1} due to missing or empty '题目' key.")
                continue

            for option in ["A", "B", "C", "D", "E"]:
                if option in q and q[option] is not None:
                    doc.add_paragraph(f"   {option}. {q[option]}")
            doc.add_paragraph("")


    # --- Add true/false questions ---
    if selected_judge:
        doc.add_heading('三、判断题', level=2)
        for i, q in enumerate(selected_judge):
            # ---  Check if '题目' key exists and is not None ---
            if '题目' in q and q['题目'] is not None:
                doc.add_paragraph(f"{q['题号']}. {q['题目']}")  # Use stored question number
            else:
                print(f"DEBUG: Skipping question {i+1} due to missing or empty '题目' key.")
                continue
            doc.add_paragraph("")


    # --- Add short answer questions ---
    if selected_short:
        doc.add_heading('四、简答题', level=2)
        for i, q in enumerate(selected_short):
            # ---  Check if '题目' key exists and is not None ---
            if '题目' in q and q['题目'] is not None:
                doc.add_paragraph(f"{q['题号']}. {q['题目']}")  # Use stored question number
            else:
                print(f"DEBUG: Skipping question {i+1} due to missing or empty '题目' key.")
                continue
            doc.add_paragraph("")

    # --- Create Answer Key Page ---
    doc.add_page_break()
    doc.add_heading('答案', level=1)

    # --- Add answers for each question type, maintaining original numbering ---
    if selected_single:
        doc.add_paragraph("一、单项选择题")
        for q in selected_single:  # Iterate through *selected* questions
            if '题号' in q and '答案' in q:  # Check for both '题号' and '答案'
                doc.add_paragraph(f"{q['题号']}. {q.get('答案', '')}")

    if selected_multi:
        doc.add_paragraph("\n二、多项选择题")  # Add a newline for separation
        for q in selected_multi:
            if '题号' in q and '答案' in q:
                doc.add_paragraph(f"{q['题号']}. {q.get('答案', '')}")

    if selected_judge:
        doc.add_paragraph("\n三、判断题")
        for q in selected_judge:
            if '题号' in q and '答案' in q:
                doc.add_paragraph(f"{q['题号']}. {q.get('答案', '')}")

    if selected_short:
        doc.add_paragraph("\n四、简答题")
        for q in selected_short:
            if '题号' in q and '题目' in q and '答案' in q:
                doc.add_paragraph(f"{q['题号']}. {q['题目']}")
                doc.add_paragraph(f"   {q.get('答案', '')}")  # Indented, no "答："
            else:
                print(f"DEBUG: Skipping short answer question due to missing data.")

    doc.save(output_path)



def select_file():
    """Opens a file selection dialog and returns the selected file path."""
    root = tk.Tk()
    root.withdraw()

    file_path = filedialog.askopenfilename(
        title="选择题库文件",
        filetypes=[("Word 文档", "*.docx"), ("所有文件", "*.*")]
    )
    return file_path

def generate_paper():
    """Main function to generate the exam paper."""
    docx_path = select_file()
    if not docx_path:
        messagebox.showinfo("提示", "未选择文件，程序退出。")
        return

    # Use a try-except block for the input dialogs as well
    try:
        num_single = simpledialog.askinteger("输入", "请输入要抽取的单选题数量：", initialvalue=10)
        if num_single is None:  # Check for cancel
            return
        num_multi = simpledialog.askinteger("输入", "请输入要抽取的多选题数量：", initialvalue=5)
        if num_multi is None:
            return
        num_judge = simpledialog.askinteger("输入", "请输入要抽取的判断题数量：", initialvalue=5)
        if num_judge is None:
            return
        num_short = simpledialog.askinteger("输入", "请输入要抽取的简答题数量：", initialvalue=2)
        if num_short is None:
            return

    except ValueError:
        messagebox.showerror("错误", "请输入有效的整数。")
        return


    output_path = filedialog.asksaveasfilename(
        title="保存试卷",
        defaultextension=".docx",
        filetypes=[("Word 文档", "*.docx"), ("所有文件", "*.*")]
    )
    if not output_path:
        messagebox.showinfo("提示", "未选择保存路径，程序退出。")
        return

    try:
        questions = extract_questions(docx_path)
        if questions:
            create_exam_paper(questions, num_single, num_multi, num_judge, num_short, output_path)
            messagebox.showinfo("提示", f"试卷已生成：{output_path}")
    except Exception as e:
        messagebox.showerror("错误", f"生成试卷时发生错误：{e}")

if __name__ == "__main__":
    generate_paper()